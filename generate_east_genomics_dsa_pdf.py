"""
Generate east_genomics_data_sharing_agreement.pdf and .docx
East Genomics branding — dual logo header, simplified front matter.
"""

import argparse
import yaml
import sys
from pathlib import Path
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether,
)
from reportlab.platypus.flowables import Flowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import Image as RLImage

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PDF = str(BASE_DIR / "east_genomics_data_sharing_agreement.pdf")
OUTPUT_DOCX = str(BASE_DIR / "east_genomics_data_sharing_agreement.docx")
LOGO_L = str(BASE_DIR / "logo_east_genomics.png")   # left logo
LOGO_R = str(BASE_DIR / "logo_east_genomics_new")   # right logo

PAGE_W, PAGE_H = A4
MARGIN = 25.4 * mm

# ---------------------------------------------------------------------------
# Colours
# ---------------------------------------------------------------------------
NHS_BLUE   = colors.HexColor("#0072C6")
HEAD_BLUE  = colors.HexColor("#2E74B5")
HEAD_DARK  = colors.HexColor("#1F4D78")
WHITE      = colors.white
LIGHT_GREY = colors.HexColor("#F2F2F2")
MID_GREY   = colors.HexColor("#D0D0D0")
BLACK      = colors.black

# ---------------------------------------------------------------------------
# AcroForm text-field flowable
# ---------------------------------------------------------------------------
class TextField(Flowable):
    def __init__(self, name, width, height=9 * mm, tooltip="", value=""):
        super().__init__()
        self.name    = name
        self.width   = width
        self.height  = height
        self.tooltip = tooltip
        self.value   = value

    def draw(self):
        self.canv.acroForm.textfield(
            name=self.name,
            tooltip=self.tooltip,
            value=self.value,
            x=0, y=0,
            width=self.width,
            height=self.height,
            borderStyle="underlined",
            borderColor=colors.HexColor("#888888"),
            fillColor=colors.Color(0, 0, 0, 0),
            textColor=BLACK,
            forceBorder=True,
            fontSize=10,
            relative=True,
        )

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
base = getSampleStyleSheet()

def ps(name, parent="Normal", **kw):
    return ParagraphStyle(name, parent=base[parent], **kw)

org_style    = ps("Org",    fontSize=13, leading=17, textColor=HEAD_BLUE,
                  spaceAfter=1, spaceBefore=0)
title_style  = ps("Title2", fontSize=18, leading=22, textColor=HEAD_BLUE,
                  spaceAfter=2, spaceBefore=0, fontName="Helvetica-Bold")
sub_style    = ps("Sub",    fontSize=12, leading=16, textColor=HEAD_DARK,
                  spaceAfter=6, spaceBefore=0)
h2_style     = ps("H2",     fontSize=13, leading=17, textColor=HEAD_BLUE,
                  spaceBefore=12, spaceAfter=4, fontName="Helvetica-Bold")
h3_style     = ps("H3",     fontSize=11, leading=15, textColor=HEAD_DARK,
                  spaceBefore=8,  spaceAfter=3, fontName="Helvetica-Bold")
body_style   = ps("Body",   fontSize=10, leading=15, spaceAfter=5)
bullet_style = ps("Bullet", fontSize=10, leading=15, leftIndent=14, spaceAfter=3)
tbl_hdr      = ps("TblHdr", fontSize=10, leading=14, textColor=WHITE,
                  fontName="Helvetica-Bold")
tbl_cell     = ps("TblCell", fontSize=10, leading=14)
footer_par   = ps("Footer", fontSize=8,  leading=11,
                  textColor=colors.HexColor("#666666"), alignment=1)

def bul(text):
    return Paragraph(f"&#8226;&nbsp;&nbsp;{text}", bullet_style)

def bold(text):
    return f"<b>{text}</b>"

# ---------------------------------------------------------------------------
# Parse arguments and load configuration
# ---------------------------------------------------------------------------
parser = argparse.ArgumentParser(description="Generate Data Sharing Agreement PDF and DOCX")
parser.add_argument("--config", default="dsa_config_community_cloud.yaml",
                    help="Path to YAML configuration file (default: dsa_config_community_cloud.yaml)")
args = parser.parse_args()

# Load YAML configuration
with open(args.config, 'r') as f:
    config = yaml.safe_load(f)

# Validate YAML schema
def validate_config(config):
    """Validate the YAML configuration has required structure."""
    required_keys = ['header', 'footer', 'metadata', 'document_details',
                     'document_control', 'sections', 'declaration']

    for key in required_keys:
        if key not in config:
            print(f"Error: Missing required top-level key '{key}' in configuration", file=sys.stderr)
            sys.exit(1)

    # Validate header
    if 'title' not in config['header']:
        print("Error: Missing 'title' in 'header' section", file=sys.stderr)
        sys.exit(1)

    # Validate footer
    if 'text' not in config['footer']:
        print("Error: Missing 'text' in 'footer' section", file=sys.stderr)
        sys.exit(1)

    # Validate metadata
    for field in ['pdf_title', 'pdf_author']:
        if field not in config['metadata']:
            print(f"Error: Missing '{field}' in 'metadata' section", file=sys.stderr)
            sys.exit(1)

    # Validate document_details
    for field in ['title', 'reference_label', 'reference_tooltip']:
        if field not in config['document_details']:
            print(f"Error: Missing '{field}' in 'document_details' section", file=sys.stderr)
            sys.exit(1)

    # Validate document_control
    for field in ['heading', 'text']:
        if field not in config['document_control']:
            print(f"Error: Missing '{field}' in 'document_control' section", file=sys.stderr)
            sys.exit(1)

    # Validate sections
    if not isinstance(config['sections'], list):
        print("Error: 'sections' must be a list", file=sys.stderr)
        sys.exit(1)

    for idx, section in enumerate(config['sections']):
        if not isinstance(section, dict):
            print(f"Error: Section {idx} must be a dict", file=sys.stderr)
            sys.exit(1)
        if 'title' not in section:
            print(f"Error: Missing 'title' in section {idx}", file=sys.stderr)
            sys.exit(1)
        if 'content' not in section:
            print(f"Error: Missing 'content' in section {idx}", file=sys.stderr)
            sys.exit(1)

    # Validate declaration
    for field in ['title', 'text', 'fields', 'footer_text']:
        if field not in config['declaration']:
            print(f"Error: Missing '{field}' in 'declaration' section", file=sys.stderr)
            sys.exit(1)

    if not isinstance(config['declaration']['fields'], list):
        print("Error: 'declaration.fields' must be a list", file=sys.stderr)
        sys.exit(1)

    for idx, field in enumerate(config['declaration']['fields']):
        if not isinstance(field, dict):
            print(f"Error: Declaration field {idx} must be a dict", file=sys.stderr)
            sys.exit(1)
        for key in ['label', 'name', 'tooltip']:
            if key not in field:
                print(f"Error: Missing '{key}' in declaration field {idx} (label: {field.get('label', 'unknown')})", file=sys.stderr)
                sys.exit(1)

validate_config(config)

# ---------------------------------------------------------------------------
# Standard table style
# ---------------------------------------------------------------------------
def std_table_style(header_rows=1):
    return TableStyle([
        ("BACKGROUND",     (0, 0), (-1, header_rows - 1), NHS_BLUE),
        ("TEXTCOLOR",      (0, 0), (-1, header_rows - 1), WHITE),
        ("FONTNAME",       (0, 0), (-1, header_rows - 1), "Helvetica-Bold"),
        ("FONTSIZE",       (0, 0), (-1, -1), 10),
        ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING",    (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 6),
        ("TOPPADDING",     (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 2),
        ("ROWBACKGROUNDS", (0, header_rows), (-1, -1), [LIGHT_GREY, WHITE]),
        ("GRID",           (0, 0), (-1, -1), 0.5, MID_GREY),
    ])

# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
def add_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    text = f"{config['footer']['text']}    |    Page {doc.page}"
    canvas.drawCentredString(PAGE_W / 2, 15 * mm, text)
    canvas.restoreState()

# ---------------------------------------------------------------------------
# Content rendering functions
# ---------------------------------------------------------------------------
def render_content_to_pdf(content_items, story):
    """Recursively render content items to PDF story."""
    for item in content_items:
        if isinstance(item, str):
            # Plain paragraph
            story.append(Paragraph(item, body_style))
        elif isinstance(item, dict):
            if 'subsection_title' in item:
                # Subsection
                story.append(Paragraph(item['subsection_title'], h3_style))
                render_content_to_pdf(item['content'], story)
            elif 'bullets' in item:
                # Bullet list
                for bullet_text in item['bullets']:
                    story.append(bul(bullet_text))
                story.append(Spacer(1, 4))

def render_content_to_docx(content_items, docx_obj):
    """Recursively render content items to DOCX document."""
    for item in content_items:
        if isinstance(item, str):
            # Plain paragraph
            para = docx_obj.add_paragraph(item)
            para.paragraph_format.space_after = Pt(5)
        elif isinstance(item, dict):
            if 'subsection_title' in item:
                # Subsection
                para = docx_obj.add_paragraph()
                run = para.add_run(item['subsection_title'])
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(31, 77, 120)
                run.bold = True
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(3)
                render_content_to_docx(item['content'], docx_obj)
            elif 'bullets' in item:
                # Bullet list
                for bullet_text in item['bullets']:
                    # Remove HTML bold tags for DOCX
                    clean_text = bullet_text.replace('<b>', '').replace('</b>', '')
                    para = docx_obj.add_paragraph(clean_text, style='List Bullet')
                    para.paragraph_format.space_after = Pt(3)

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
CONTENT_W = PAGE_W - 2 * MARGIN
PAD       = 6 * mm

doc = SimpleDocTemplate(
    OUTPUT_PDF,
    pagesize=A4,
    leftMargin=MARGIN, rightMargin=MARGIN,
    topMargin=MARGIN,  bottomMargin=22 * mm,
    title=config['metadata']['pdf_title'],
    author=config['metadata']['pdf_author'],
)

story = []

# ── Dual-logo header ──────────────────────────────────────────────────────────
# Both logos rendered at the same fixed height; widths derived from aspect ratio.
# Left logo: 600x300 (2:1)  →  width = 2 × H
# Right logo: 353x353 (1:1) →  width = 1 × H
LOGO_H   = 22 * mm
LOGO_L_W = 44 * mm   # 2:1 aspect × 22mm
LOGO_R_W = 56.3 * mm * 0.8  # 852:333 aspect × 22mm, scaled to 80%
TEXT_W   = CONTENT_W - LOGO_L_W - LOGO_R_W - 4 * mm

try:
    logo_left  = RLImage(LOGO_L, width=LOGO_L_W, height=LOGO_H)
except Exception:
    logo_left  = Spacer(LOGO_L_W, LOGO_H)

try:
    logo_right = RLImage(LOGO_R, width=LOGO_R_W, height=LOGO_H * 0.8)
except Exception:
    logo_right = Spacer(LOGO_R_W, LOGO_H)

_title_centre = ParagraphStyle("TitleCentre", parent=title_style, alignment=1)
title_block = Table([
    [Paragraph(config['header']['title'], _title_centre)],
], colWidths=[TEXT_W])

header_table = Table(
    [[logo_left, title_block, logo_right]],
    colWidths=[LOGO_L_W, TEXT_W, LOGO_R_W],
)
header_table.setStyle(TableStyle([
    ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
    ("ALIGN",         (2, 0), (2, 0),   "RIGHT"),
    ("LEFTPADDING",   (0, 0), (-1, -1), 0),
    ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
    ("TOPPADDING",    (0, 0), (-1, -1), 0),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ("LINEBELOW",     (0, 0), (-1, 0),  2, NHS_BLUE),
]))
story.append(header_table)
story.append(Spacer(1, 16))

# ── Document details ──────────────────────────────────────────────────────────
COL_LABEL = 60 * mm
COL_VALUE = CONTENT_W - COL_LABEL
FW_VALUE  = COL_VALUE - 2 * PAD

story.append(Paragraph("Title", h3_style))
story.append(Paragraph(config['document_details']['title'], body_style))
story.append(Spacer(1, 10))

story.append(Paragraph(config['document_details']['reference_label'], h3_style))
story.append(TextField(name="det_ref", width=CONTENT_W, height=9 * mm,
                       tooltip=config['document_details']['reference_tooltip']))
story.append(Spacer(1, 16))

# ── Document control ──────────────────────────────────────────────────────────
story.append(Paragraph(config['document_control']['heading'], h3_style))
story.append(Paragraph(config['document_control']['text'], body_style))

story.append(PageBreak())

# ============================================================================
# Main agreement content
# ============================================================================

for section in config['sections']:
    story.append(Paragraph(section['title'], h2_style))
    render_content_to_pdf(section['content'], story)

# ── Participant Declaration ──────────────────────────────────────────────────
story.append(Paragraph(config['declaration']['title'], h2_style))
story.append(Paragraph(config['declaration']['text'], body_style))
story.append(Spacer(1, 6))

LABEL_W   = 48 * mm
FIELD_COL = CONTENT_W - LABEL_W

decl_data = [
    [Paragraph(bold("Field"), tbl_hdr),
     Paragraph(bold("Details"), tbl_hdr)],
]
for field in config['declaration']['fields']:
    decl_data.append([
        Paragraph(field['label'], tbl_cell),
        TextField(name=field['name'], width=FIELD_COL - 2 * PAD, tooltip=field['tooltip']),
    ])

row_heights = [10 * mm] + [12 * mm] * len(config['declaration']['fields'])
decl_table = Table(
    decl_data,
    colWidths=[LABEL_W, FIELD_COL],
    rowHeights=row_heights,
    repeatRows=1,
)
decl_table.setStyle(TableStyle([
    ("BACKGROUND",    (0, 0), (-1, 0),  NHS_BLUE),
    ("TEXTCOLOR",     (0, 0), (-1, 0),  WHITE),
    ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
    ("ROWBACKGROUNDS",(0, 1), (-1, -1), [LIGHT_GREY, WHITE]),
    ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
    ("LEFTPADDING",   (0, 0), (-1, -1), PAD),
    ("RIGHTPADDING",  (0, 0), (-1, -1), PAD),
    ("TOPPADDING",    (0, 0), (-1, -1), 2),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ("GRID",          (0, 0), (-1, -1), 0.5, MID_GREY),
]))

story.append(decl_table)
story.append(Spacer(1, 10))
story.append(HRFlowable(width="100%", thickness=0.5, color=MID_GREY))
story.append(Spacer(1, 4))
story.append(Paragraph(config['declaration']['footer_text'], footer_par))

doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
print(f"PDF written to: {OUTPUT_PDF}")

# ============================================================================
# Generate DOCX version
# ============================================================================

def add_text_field(paragraph, label):
    """Add a form text field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' FORMTEXT '
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

    # Add underline placeholder
    run = paragraph.add_run('_' * 60)
    run.font.color.rgb = RGBColor(128, 128, 128)

docx = Document()

# Set margins
sections = docx.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.87)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Header with logos and title
header = docx.sections[0].header
header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
header_table.rows[0].cells[0].width = Inches(2.2)
header_table.rows[0].cells[1].width = Inches(2.1)
header_table.rows[0].cells[2].width = Inches(2.2)

# Add left logo
try:
    para = header_table.rows[0].cells[0].paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()
    run.add_picture(LOGO_L, width=Inches(1.73))
except:
    pass

# Add title
para = header_table.rows[0].cells[1].paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run(config['header']['title'])
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(46, 116, 181)
run.bold = True

# Add right logo
try:
    para = header_table.rows[0].cells[2].paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    run.add_picture(LOGO_R, width=Inches(1.8))
except:
    pass

# Add horizontal line under header
para = header.add_paragraph()
para.paragraph_format.space_after = Pt(12)
run = para.add_run()

# Document details
para = docx.add_paragraph()
run = para.add_run("Title")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(31, 77, 120)
run.bold = True
para.paragraph_format.space_before = Pt(8)
para.paragraph_format.space_after = Pt(3)

para = docx.add_paragraph(config['document_details']['title'])
para.paragraph_format.space_after = Pt(10)

para = docx.add_paragraph()
run = para.add_run(config['document_details']['reference_label'])
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(31, 77, 120)
run.bold = True
para.paragraph_format.space_before = Pt(8)
para.paragraph_format.space_after = Pt(3)

para = docx.add_paragraph()
add_text_field(para, config['document_details']['reference_tooltip'])
para.paragraph_format.space_after = Pt(16)

# Document control
para = docx.add_paragraph()
run = para.add_run(config['document_control']['heading'])
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(31, 77, 120)
run.bold = True
para.paragraph_format.space_before = Pt(8)
para.paragraph_format.space_after = Pt(4)

para = docx.add_paragraph(config['document_control']['text'])
para.paragraph_format.space_after = Pt(5)

docx.add_page_break()

# Main content sections
for section in config['sections']:
    # Add section title
    para = docx.add_paragraph()
    run = para.add_run(section['title'])
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(46, 116, 181)
    run.bold = True
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)

    # Render section content
    render_content_to_docx(section['content'], docx)

# Participant Declaration
para = docx.add_paragraph()
run = para.add_run(config['declaration']['title'])
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(46, 116, 181)
run.bold = True
para.paragraph_format.space_before = Pt(12)
para.paragraph_format.space_after = Pt(4)

para = docx.add_paragraph(config['declaration']['text'])
para.paragraph_format.space_after = Pt(6)

# Declaration table
num_fields = len(config['declaration']['fields'])
table = docx.add_table(rows=num_fields + 1, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = "Field"
header_cells[1].text = "Details"
for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    # Set background color to NHS Blue
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '0072C6')
    cell._element.get_or_add_tcPr().append(shading_elm)

# Data rows
for idx, field in enumerate(config['declaration']['fields'], start=1):
    table.rows[idx].cells[0].text = field['label']
    para = table.rows[idx].cells[1].paragraphs[0]
    add_text_field(para, field['label'])

para = docx.add_paragraph()
para.paragraph_format.space_before = Pt(10)
para.paragraph_format.space_after = Pt(4)

para = docx.add_paragraph(config['declaration']['footer_text'])
para.runs[0].font.size = Pt(8)
para.runs[0].font.color.rgb = RGBColor(102, 102, 102)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add footer
footer = docx.sections[0].footer
para = footer.paragraphs[0]
para.text = config['footer']['text']
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.runs[0].font.size = Pt(8)
para.runs[0].font.color.rgb = RGBColor(102, 102, 102)

docx.save(OUTPUT_DOCX)
print(f"DOCX written to: {OUTPUT_DOCX}")
